"""
PDF Extractor Tool
-----------------
Author: Haytham Abo Abdallah
Created: February 2025
Purpose: Created for internal company use

Description:
This tool extracts PDF hyperlinks from Word documents and processes them by:
1. Extracting all PDF links from a selected Word document
2. Copying the linked PDFs from a central folder to a destination folder
3. Adding a cover page with "المستند رقم" (Document Number) to each PDF
4. Renaming the PDFs with a consistent numbering system
5. Creating an Excel report with clickable links and document information

The tool is designed to help organize and index PDF documents for easier reference 
and management of company documentation.
"""

import os
import shutil
import docx
import urllib.parse
import openpyxl
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from tkinter import filedialog, messagebox, simpledialog
import threading
import logging
import re
from datetime import datetime
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from PyPDF2 import PdfReader, PdfWriter
import io
import sys
import platform
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
try:
    from bidi.algorithm import get_display
    import arabic_reshaper
    ARABIC_SUPPORT = True
except ImportError:
    ARABIC_SUPPORT = False

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

class PDFExtractorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("PDF Extractor Tool - by Haytham Abo Abdallah")
        self.root.geometry("650x600")
        self.root.resizable(False, False)
        
        # Version info
        self.version = "1.0"
        self.created_date = "February 2025"

        # Apply a dark mode theme
        self.style = ttk.Style("superhero")  

        # Variables to store user selections
        self.word_file_path = ttk.StringVar()
        self.central_pdf_folder = ttk.StringVar()
        self.destination_folder = ttk.StringVar()
        self.status_var = ttk.StringVar(value="Ready")
        self.enable_renaming = ttk.BooleanVar(value=True)
        self.add_cover_page = ttk.BooleanVar(value=True)
        
        # Store hyperlinks for functionality
        self.current_hyperlinks = []

        # Try to register Arabic fonts
        self.register_arabic_fonts()
        
        # Register Times-Roman as a fallback
        try:
            pdfmetrics.registerFont(UnicodeCIDFont('Times-Roman'))
        except:
            pass

        # UI Components
        self.create_widgets()
        
    def register_arabic_fonts(self):
        """Try to register Arabic fonts from system locations"""
        # List of potential Arabic font locations based on OS
        font_paths = []
        
        system = platform.system()
        if system == 'Windows':
            windows_font_dir = os.path.join(os.environ.get('SystemRoot', 'C:\\Windows'), 'Fonts')
            font_paths = [
                os.path.join(windows_font_dir, 'arial.ttf'),
                os.path.join(windows_font_dir, 'tahoma.ttf'),
                os.path.join(windows_font_dir, 'times.ttf'),
                os.path.join(windows_font_dir, 'arabtype.ttf'),
                os.path.join(windows_font_dir, 'simpo.ttf'),
                os.path.join(windows_font_dir, 'simpfxo.ttf'),
            ]
        elif system == 'Darwin':  # macOS
            font_paths = [
                '/Library/Fonts/Arial.ttf',
                '/Library/Fonts/ArialHB.ttc',
                '/System/Library/Fonts/Helvetica.ttc',
            ]
        else:  # Linux
            font_paths = [
                '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
                '/usr/share/fonts/truetype/freefont/FreeSans.ttf',
                '/usr/share/fonts/TTF/DejaVuSans.ttf',
            ]
        
        # Try to register fonts until one works
        for font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    pdfmetrics.registerFont(TTFont('Arabic', font_path))
                    logging.info(f"Registered font from {font_path}")
                    return True
                except Exception as e:
                    logging.warning(f"Failed to register font {font_path}: {e}")
                    continue
        
        logging.warning("Could not register any Arabic fonts")
        return False



    def create_widgets(self):
        """Create modern UI components"""

        # Title Label
        ttk.Label(self.root, text="📄 PDF Extractor Tool", font=("Arial", 18, "bold"), bootstyle=PRIMARY).pack(pady=10)
        
        # Description
        description = "Extract, rename, and organize PDF files from Word documents"
        ttk.Label(self.root, text=description, font=("Arial", 10), bootstyle=SECONDARY).pack(pady=0)

        # Select Word Document
        self.create_file_selection("Select Word File:", self.word_file_path, self.select_word_file)

        # Select Central PDF Folder
        self.create_folder_selection("Central PDF Folder:", self.central_pdf_folder, self.select_central_pdf_folder)

        # Select Destination Folder
        self.create_folder_selection("Destination Folder:", self.destination_folder, self.select_destination_folder)

        # Options Frame
        options_frame = ttk.Frame(self.root)
        options_frame.pack(pady=5, fill="x", padx=20)
        
        # Renaming Options
        ttk.Label(options_frame, text="PDF Options:", font=("Arial", 10, "bold"), bootstyle=LIGHT).pack(anchor="w")
        ttk.Checkbutton(options_frame, text="Enable Renaming (Prefix with 'المستند رقم')", 
                      variable=self.enable_renaming, bootstyle="round-toggle").pack(anchor="w", pady=2)
        ttk.Checkbutton(options_frame, text="Add Cover Page", 
                      variable=self.add_cover_page, bootstyle="round-toggle").pack(anchor="w", pady=2)

        # Progress Bar
        self.progress = ttk.Progressbar(self.root, mode="determinate", bootstyle=INFO, length=500)
        self.progress.pack(pady=15)

        # Status Label
        ttk.Label(self.root, textvariable=self.status_var, font=("Arial", 10), bootstyle="info").pack(pady=5)

        # Process Button
        ttk.Button(self.root, text="🚀 Process Word File", command=self.start_processing, bootstyle=SUCCESS, width=25).pack(pady=10)
        
        # Author info
        info_frame = ttk.Frame(self.root)
        info_frame.pack(side="bottom", fill="x", padx=20, pady=10)
        
        author_text = f"Created by Haytham Abo Abdallah | v{self.version} | {self.created_date}"
        ttk.Label(info_frame, text=author_text, font=("Arial", 8), 
                bootstyle=SECONDARY).pack(side="right")

    def create_file_selection(self, label_text, variable, command):
        """Create a file selection UI row"""
        frame = ttk.Frame(self.root)
        frame.pack(pady=5, fill="x", padx=20)
        ttk.Label(frame, text=label_text, font=("Arial", 10, "bold"), bootstyle=LIGHT).pack(anchor="w")
        entry = ttk.Entry(frame, textvariable=variable, width=50, bootstyle="secondary")
        entry.pack(side="left", padx=5)
        ttk.Button(frame, text="Browse", command=command, bootstyle=PRIMARY).pack(side="right")

    def create_folder_selection(self, label_text, variable, command):
        """Create a folder selection UI row"""
        frame = ttk.Frame(self.root)
        frame.pack(pady=5, fill="x", padx=20)
        ttk.Label(frame, text=label_text, font=("Arial", 10, "bold"), bootstyle=LIGHT).pack(anchor="w")
        entry = ttk.Entry(frame, textvariable=variable, width=50, bootstyle="secondary")
        entry.pack(side="left", padx=5)
        ttk.Button(frame, text="Browse", command=command, bootstyle=PRIMARY).pack(side="right")

    def select_word_file(self):
        """Open file dialog to select the Word document"""
        file_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
        if file_path:
            self.word_file_path.set(file_path)
            # Extract hyperlinks when file is selected
            self.extract_links_from_word()

    def select_central_pdf_folder(self):
        """Open folder dialog to select the central PDF folder"""
        folder_path = filedialog.askdirectory()
        if folder_path:
            self.central_pdf_folder.set(folder_path)

    def select_destination_folder(self):
        """Open folder dialog to select the destination folder"""
        folder_path = filedialog.askdirectory()
        if folder_path:
            self.destination_folder.set(folder_path)

    def extract_links_from_word(self):
        """Extract links from the selected Word file"""
        word_file_path = self.word_file_path.get()
        if not word_file_path:
            return

        try:
            self.status_var.set("Extracting links from Word file...")
            self.root.update_idletasks()
            
            doc = docx.Document(word_file_path)
            self.current_hyperlinks = self.extract_hyperlinks(doc)
            
            if self.current_hyperlinks:
                self.status_var.set(f"Found {len(self.current_hyperlinks)} PDF links in document")
            else:
                self.status_var.set("No PDF links found in the document")
                
        except Exception as e:
            self.status_var.set("Error extracting links")
            messagebox.showerror("Error", f"❌ Failed to extract links: {e}")

    def start_processing(self):
        """Start the processing in a separate thread to keep UI responsive"""
        thread = threading.Thread(target=self.process_word_file, daemon=True)
        thread.start()

    def process_word_file(self):
        """Process the selected Word document"""
        word_file_path = self.word_file_path.get()
        central_pdf_folder = self.central_pdf_folder.get()
        destination_folder = self.destination_folder.get()

        if not word_file_path or not central_pdf_folder or not destination_folder:
            messagebox.showerror("Error", "⚠️ Please select all required paths.")
            return

        try:
            doc = docx.Document(word_file_path)
        except Exception as e:
            messagebox.showerror("Error", f"❌ Failed to load Word document: {e}")
            return

        self.progress.start(10)
        self.status_var.set("Extracting hyperlinks...")
        self.root.update_idletasks()

        hyperlinks = self.extract_hyperlinks(doc)
        if not hyperlinks:
            self.progress.stop()
            self.status_var.set("No PDF hyperlinks found")
            messagebox.showinfo("Info", "ℹ️ No PDF hyperlinks found in the document.")
            return

        self.status_var.set(f"Found {len(hyperlinks)} PDF links")
        
        # Check which PDFs exist
        existing_pdfs = []
        missing_pdfs = []
        
        for link_text, pdf_filename in hyperlinks:
            pdf_path = os.path.join(central_pdf_folder, os.path.basename(pdf_filename))
            if os.path.exists(pdf_path):
                existing_pdfs.append((link_text, pdf_filename))
            else:
                missing_pdfs.append((link_text, pdf_filename))
        
        if missing_pdfs:
            result = messagebox.askyesno(
                "Missing PDFs", 
                f"⚠️ {len(missing_pdfs)} PDF files were not found in the central folder.\nDo you want to continue with the {len(existing_pdfs)} available PDFs only?"
            )
            if not result:
                self.progress.stop()
                self.status_var.set("Operation cancelled")
                return
        
        self.status_var.set("Processing PDF files...")
        
        # Process PDFs first, then create Excel with correct links
        if existing_pdfs:
            if self.enable_renaming.get() or self.add_cover_page.get():
                self.process_pdfs(existing_pdfs, central_pdf_folder, destination_folder)
            else:
                self.copy_files_parallel(existing_pdfs, central_pdf_folder, destination_folder)
        else:
            self.status_var.set("No PDFs to process")

        self.status_var.set("Copying Word file...")
        self.copy_word_file(word_file_path, destination_folder)
        
        self.status_var.set("Saving links to Excel...")
        xlsx_file_path = os.path.join(destination_folder, "extracted_links.xlsx")
        self.save_links_to_xlsx(hyperlinks, existing_pdfs, missing_pdfs, xlsx_file_path, destination_folder)

        self.progress.stop()
        self.progress['value'] = 100
        self.status_var.set("Processing completed!")
        
        # Open the destination folder to show the results
        try:
            if platform.system() == "Windows":
                os.startfile(destination_folder)
            elif platform.system() == "Darwin":  # macOS
                import subprocess
                subprocess.call(["open", destination_folder])
            else:  # Linux
                import subprocess
                subprocess.call(["xdg-open", destination_folder])
        except:
            pass
            
        messagebox.showinfo("Success", f"✅ Processing completed! \n📂 XLSX saved at: {xlsx_file_path}")



    def extract_hyperlinks(self, doc):
        """Extracts PDF hyperlinks from a Word document."""
        hyperlinks = []
        for para in doc.paragraphs:
            for hyperlink in para._element.xpath(".//w:hyperlink"):
                link_text = "".join(t.text for t in hyperlink.xpath(".//w:r/w:t")) or "Unnamed Link"
                r_id = hyperlink.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
                if r_id and r_id in doc.part.rels:
                    url = os.path.normpath(urllib.parse.unquote(doc.part.rels[r_id].target_ref.strip()))
                    if url.lower().endswith(".pdf"):
                        hyperlinks.append((link_text, url))
        return hyperlinks

    def save_links_to_xlsx(self, all_hyperlinks, existing_pdfs, missing_pdfs, xlsx_file_path, destination_folder):
        """Saves extracted hyperlinks to an XLSX file with status and document numbers."""
        workbook = openpyxl.Workbook()
        
        # Create All Links sheet
        sheet1 = workbook.active
        sheet1.title = "All Links"
        
        # Set column headers
        sheet1.append(["Link Text", "PDF Link", "Document Number", "Status"])
        
        # Create a dictionary for quick lookup of status
        existing_dict = {url: "Found" for _, url in existing_pdfs}
        missing_dict = {url: "Missing" for _, url in missing_pdfs}
        
        # Track unique filenames to prevent duplicate document numbers
        unique_filenames = {}
        doc_number_counter = 1
        
        # First pass: identify unique filenames
        for _, url in existing_pdfs:
            filename = os.path.basename(url)
            if filename not in unique_filenames:
                unique_filenames[filename] = doc_number_counter
                doc_number_counter += 1
                
        # Create dictionaries for document numbers and renamed files
        doc_number_dict = {}
        renamed_files = {}
        
        for _, url in existing_pdfs:
            filename = os.path.basename(url)
            doc_number = unique_filenames[filename]
            doc_number_dict[url] = f"المستند رقم {doc_number:03d}"
            
            # Create the renamed filename as used in process_pdfs
            new_name = f"المستند رقم {doc_number:03d} - {filename}"
            renamed_files[url] = new_name
        
        # Add all links with status and document number
        for link_text, url in all_hyperlinks:
            status = existing_dict.get(url, missing_dict.get(url, "Unknown"))
            doc_number = doc_number_dict.get(url, "")
            filename = os.path.basename(url)
            
            # Add the row
            sheet1.append([link_text, filename, doc_number, status])
            
            # Make the URL in the second column clickable to the destination file
            if status == "Found":
                cell = sheet1.cell(row=sheet1.max_row, column=2)
                new_filename = renamed_files.get(url, filename)
                # Create local file:// hyperlink to the destination folder
                local_path = os.path.abspath(os.path.join(destination_folder, new_filename))
                cell.hyperlink = f"file:///{local_path.replace('\\', '/')}"
                cell.style = "Hyperlink"
            
        # Add conditional formatting for status column
        from openpyxl.styles import PatternFill
        green_fill = PatternFill(start_color="92D050", end_color="92D050", fill_type="solid")
        red_fill = PatternFill(start_color="FF6B6B", end_color="FF6B6B", fill_type="solid")
        
        # Apply color to all cells in each row
        for row in range(2, sheet1.max_row + 1):
            status = sheet1.cell(row=row, column=4).value
            if status == "Found":
                for col in range(1, 5):  # Color all four columns
                    cell = sheet1.cell(row=row, column=col)
                    if col == 4:  # Status column
                        cell.fill = green_fill
                    else:
                        cell.fill = PatternFill(start_color="E6FFE6", end_color="E6FFE6", fill_type="solid")  # Light green
            elif status == "Missing":
                sheet1.cell(row=row, column=4).fill = red_fill
                
        # Auto-adjust column widths
        for column in sheet1.columns:
            max_length = 0
            column_letter = openpyxl.utils.get_column_letter(column[0].column)
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2) * 1.2
            sheet1.column_dimensions[column_letter].width = adjusted_width
                
        workbook.save(xlsx_file_path)

    def create_cover_page(self, doc_number, link_text, pdf_filename):
        """Create a PDF cover page with 'المستند رقم' and the document number"""
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        
        # Format the document number
        formatted_number = f"{doc_number}"
        
        # Text to display (المستند رقم + number)
        arabic_text = f"المستند رقم {formatted_number}"
        
        # Try with Arabic reshaper and bidi algorithm if available
        if ARABIC_SUPPORT:
            try:
                # Reshape the Arabic text
                reshaped_text = arabic_reshaper.reshape(arabic_text)
                # Apply bidirectional algorithm
                bidi_text = get_display(reshaped_text)
                
                # Use a registered font
                c.setFont("Arabic", 24)
                # Center the text on the page
                c.drawCentredString(width/2, height/2, bidi_text)
                
                logging.info("Cover page created with Arabic reshaper")
            except Exception as e:
                logging.error(f"Error with Arabic reshaper: {e}")
                # Fall back to direct method
                c.setFont("Helvetica-Bold", 36)
                c.drawCentredString(width/2, height/2, formatted_number)
        else:
            # If bidi/reshaper not available, try with basic RTL text
            try:
                # Use Times-Roman which might have better Arabic support
                c.setFont("Times-Roman", 24)
                # Try to use raw Arabic text - may not work in all systems
                c.drawCentredString(width/2, height/2, arabic_text)
            except:
                # Last resort fallback
                c.setFont("Helvetica-Bold", 36)
                c.drawCentredString(width/2, height/2, formatted_number)
        
        c.save()
        buffer.seek(0)
        return buffer

    def process_pdfs(self, existing_pdfs, central_pdf_folder, destination_folder):
        """Process PDFs - rename and add cover page as needed"""
        os.makedirs(destination_folder, exist_ok=True)
        
        # Track unique filenames to prevent duplicate document numbers
        unique_filenames = {}
        doc_number_counter = 1
        
        # First identify unique filenames and assign document numbers
        for _, pdf_filename in existing_pdfs:
            filename = os.path.basename(pdf_filename)
            if filename not in unique_filenames:
                unique_filenames[filename] = doc_number_counter
                doc_number_counter += 1
        
        total_files = len(existing_pdfs)
        processed_count = 0
        
        for link_text, pdf_filename in existing_pdfs:
            filename = os.path.basename(pdf_filename)
            doc_number = unique_filenames[filename]
            source_path = os.path.join(central_pdf_folder, filename)
            
            # Determine the new filename with updated naming format
            if self.enable_renaming.get():
                new_name = f"المستند رقم {doc_number:03d} - {filename}"
            else:
                new_name = filename
                
            dest_path = os.path.join(destination_folder, new_name)
            
            processed_count += 1
            
            # Update progress
            self.progress['value'] = (processed_count / total_files) * 100
            self.status_var.set(f"Processing {processed_count}/{total_files}: {new_name}")
            self.root.update_idletasks()
            
            try:
                if self.add_cover_page.get():
                    # Create and add cover page
                    cover_buffer = self.create_cover_page(doc_number, link_text, pdf_filename)
                    
                    # Merge cover page with original PDF
                    pdf_writer = PdfWriter()
                    
                    # Add cover page
                    cover_reader = PdfReader(cover_buffer)
                    pdf_writer.add_page(cover_reader.pages[0])
                    
                    # Add original PDF pages
                    original_reader = PdfReader(source_path)
                    for page in original_reader.pages:
                        pdf_writer.add_page(page)
                    
                    # Write to destination
                    with open(dest_path, 'wb') as output_file:
                        pdf_writer.write(output_file)
                else:
                    # Just copy the file if no cover page needed
                    shutil.copy(source_path, dest_path)
            except Exception as e:
                logging.error(f"Error processing {source_path}: {e}")
                # Try to copy the original as fallback
                try:
                    shutil.copy(source_path, dest_path)
                except:
                    pass

    def copy_files_parallel(self, hyperlinks, central_pdf_folder, destination_folder):
        """Copies PDFs in parallel using multi-threading"""
        os.makedirs(destination_folder, exist_ok=True)
        threads = []
        for _, pdf_filename in hyperlinks:
            pdf_path = os.path.join(central_pdf_folder, os.path.basename(pdf_filename))
            if os.path.exists(pdf_path):
                thread = threading.Thread(target=self.copy_file, args=(pdf_path, destination_folder))
                threads.append(thread)
                thread.start()

        for thread in threads:
            thread.join()

    def copy_file(self, source, destination):
        """Copy a single file"""
        try:
            shutil.copy(source, destination)
        except Exception as e:
            logging.error(f"Error copying {source}: {e}")

    def copy_word_file(self, word_file_path, destination_folder):
        """Copies the selected Word document to the destination folder."""
        shutil.copy(word_file_path, destination_folder)


if __name__ == "__main__":
    print("PDF Extractor Tool v1.0")
    print("Author: Haytham Abo Abdallah")
    print("=============================")
    print("Description: This tool extracts PDF hyperlinks from Word documents,")
    print("processes them by copying to a destination folder, adding cover pages,")
    print("renaming with a consistent numbering system, and creating an Excel report.")
    print("Created for internal company use.")
    print("=============================")
    
    # Check for Arabic support packages
    if ARABIC_SUPPORT:
        print("✓ Arabic text support is enabled (bidi and arabic_reshaper packages found)")
    else:
        print("! For proper Arabic text rendering, please install these packages:")
        print("  pip install arabic-reshaper python-bidi")
    
    # Check for font support
    if platform.system() == 'Windows':
        print("\nChecking for Arabic-compatible fonts:")
        windows_font_dir = os.path.join(os.environ.get('SystemRoot', 'C:\\Windows'), 'Fonts')
        arabic_fonts = [
            'arial.ttf',
            'tahoma.ttf', 
            'arabtype.ttf',  # Arabic Typesetting
            'simpo.ttf'      # Simplified Arabic
        ]
        for font in arabic_fonts:
            font_path = os.path.join(windows_font_dir, font)
            if os.path.exists(font_path):
                print(f"  ✓ Found font: {font}")
            else:
                print(f"  ✗ Missing font: {font}")
    
    print("\nStarting application...\n")
    
    try:
        root = ttk.Window(themename="superhero")
        app = PDFExtractorApp(root)
        root.mainloop()
    except Exception as e:
        print(f"Error starting application: {e}")
        input("Press Enter to exit...")